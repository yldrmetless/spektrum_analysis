"""
Gelişmiş Dosya İşlemleri Utilities
Çoklu format desteği ve gelişmiş import/export özellikleri
"""

import os
import json
import xml.etree.ElementTree as ET
from pathlib import Path
import pandas as pd
import numpy as np
from tkinter import filedialog, messagebox
import zipfile
from datetime import datetime
import tempfile

try:
    from docx import Document
    from docx.shared import Inches
    PYTHON_DOCX_AVAILABLE = True
except ImportError:
    PYTHON_DOCX_AVAILABLE = False

try:
    from openpyxl.styles import Font, Alignment, PatternFill
    from openpyxl.chart import ScatterChart, Reference, Series
    OPENPYXL_ADVANCED_AVAILABLE = True
except ImportError:
    OPENPYXL_ADVANCED_AVAILABLE = False

class AdvancedFileUtils:
    """Gelişmiş dosya işlemleri sınıfı"""
    
    @staticmethod
    def get_supported_formats():
        """Desteklenen dosya formatlarını döndürür"""
        return {
            'import': {
                'excel': ['.xlsx', '.xls'],
                'csv': ['.csv', '.txt'],
                'json': ['.json'],
                'xml': ['.xml'],
                'zip': ['.zip']
            },
            'export': {
                'excel': ['.xlsx'],
                'csv': ['.csv'],
                'json': ['.json'],
                'xml': ['.xml'],
                'word': ['.docx'] if PYTHON_DOCX_AVAILABLE else [],
                'pdf': ['.pdf'],
                'zip': ['.zip']
            }
        }
    
    @staticmethod
    def import_data_advanced(file_types=None):
        """
        Gelişmiş veri import işlemi
        
        Args:
            file_types: Desteklenecek dosya türleri
            
        Returns:
            dict: {
                'success': bool,
                'data': DataFrame veya dict,
                'metadata': dict,
                'file_info': dict
            }
        """
        # Dosya türlerini belirle
        if file_types is None:
            filetypes = [
                ("Excel Dosyaları", "*.xlsx *.xls"),
                ("CSV Dosyaları", "*.csv *.txt"), 
                ("JSON Dosyaları", "*.json"),
                ("XML Dosyaları", "*.xml"),
                ("Zip Arşivleri", "*.zip"),
                ("Tüm Dosyalar", "*.*")
            ]
        else:
            filetypes = file_types
        
        # Dosya seçim dialonu
        file_path = filedialog.askopenfilename(
            title="Gelişmiş Veri Import",
            filetypes=filetypes
        )
        
        if not file_path:
            return {'success': False, 'message': 'Dosya seçilmedi'}
        
        try:
            # Dosya uzantısını belirle
            ext = Path(file_path).suffix.lower()
            
            # Import işlemini gerçekleştir
            if ext in ['.xlsx', '.xls']:
                return AdvancedFileUtils._import_excel_advanced(file_path)
            elif ext in ['.csv', '.txt']:
                return AdvancedFileUtils._import_csv_advanced(file_path)
            elif ext == '.json':
                return AdvancedFileUtils._import_json(file_path)
            elif ext == '.xml':
                return AdvancedFileUtils._import_xml(file_path)
            elif ext == '.zip':
                return AdvancedFileUtils._import_zip(file_path)
            else:
                return {'success': False, 'message': f'Desteklenmeyen format: {ext}'}
                
        except Exception as e:
            return {
                'success': False, 
                'message': f'Import hatası: {str(e)}',
                'error': str(e)
            }
    
    @staticmethod
    def _import_excel_advanced(file_path):
        """Gelişmiş Excel import"""
        # Tüm sheet'leri al
        excel_file = pd.ExcelFile(file_path)
        sheets_data = {}
        
        for sheet_name in excel_file.sheet_names:
            df = pd.read_excel(file_path, sheet_name=sheet_name)
            sheets_data[sheet_name] = df
        
        # Ana veriyi belirle (en büyük sheet)
        main_sheet = max(sheets_data.keys(), key=lambda x: len(sheets_data[x]))
        main_data = sheets_data[main_sheet]
        
        return {
            'success': True,
            'data': main_data,
            'all_sheets': sheets_data,
            'metadata': {
                'sheet_count': len(sheets_data),
                'main_sheet': main_sheet,
                'total_rows': sum(len(df) for df in sheets_data.values())
            },
            'file_info': {
                'path': file_path,
                'name': Path(file_path).name,
                'size': Path(file_path).stat().st_size,
                'format': 'Excel'
            }
        }
    
    @staticmethod
    def _import_csv_advanced(file_path):
        """Gelişmiş CSV import"""
        # Encoding detection
        encodings = ['utf-8', 'utf-8-sig', 'latin-1', 'cp1254', 'iso-8859-9']
        df = None
        used_encoding = None
        
        for encoding in encodings:
            try:
                # Farklı separator'ları dene
                separators = [',', ';', '\t', '|']
                for sep in separators:
                    try:
                        test_df = pd.read_csv(file_path, encoding=encoding, sep=sep, nrows=5)
                        if len(test_df.columns) > 1:  # Başarılı ayrıştırma
                            df = pd.read_csv(file_path, encoding=encoding, sep=sep)
                            used_encoding = encoding
                            break
                    except:
                        continue
                if df is not None:
                    break
            except:
                continue
        
        if df is None:
            return {'success': False, 'message': 'CSV dosyası okunamadı'}
        
        return {
            'success': True,
            'data': df,
            'metadata': {
                'encoding': used_encoding,
                'rows': len(df),
                'columns': len(df.columns)
            },
            'file_info': {
                'path': file_path,
                'name': Path(file_path).name,
                'size': Path(file_path).stat().st_size,
                'format': 'CSV'
            }
        }
    
    @staticmethod
    def _import_json(file_path):
        """JSON import"""
        with open(file_path, 'r', encoding='utf-8') as f:
            data = json.load(f)
        
        # JSON'u DataFrame'e çevirmeye çalış
        try:
            if isinstance(data, list):
                df = pd.DataFrame(data)
            elif isinstance(data, dict) and 'data' in data:
                df = pd.DataFrame(data['data'])
            else:
                df = pd.json_normalize(data)
        except:
            df = None
        
        return {
            'success': True,
            'data': df if df is not None else data,
            'raw_data': data,
            'metadata': {
                'type': type(data).__name__,
                'size': len(str(data))
            },
            'file_info': {
                'path': file_path,
                'name': Path(file_path).name,
                'size': Path(file_path).stat().st_size,
                'format': 'JSON'
            }
        }
    
    @staticmethod
    def _import_xml(file_path):
        """XML import"""
        tree = ET.parse(file_path)
        root = tree.getroot()
        
        # XML'i dict'e çevir
        def xml_to_dict(element):
            result = {}
            for child in element:
                if len(child) == 0:
                    result[child.tag] = child.text
                else:
                    result[child.tag] = xml_to_dict(child)
            return result
        
        data = xml_to_dict(root)
        
        return {
            'success': True,
            'data': data,
            'metadata': {
                'root_tag': root.tag,
                'element_count': len(list(root.iter()))
            },
            'file_info': {
                'path': file_path,
                'name': Path(file_path).name,
                'size': Path(file_path).stat().st_size,
                'format': 'XML'
            }
        }
    
    @staticmethod
    def _import_zip(file_path):
        """ZIP arşivi import"""
        extracted_files = []
        temp_dir = tempfile.mkdtemp()
        
        with zipfile.ZipFile(file_path, 'r') as zip_ref:
            zip_ref.extractall(temp_dir)
            extracted_files = zip_ref.namelist()
        
        # Excel/CSV dosyalarını ara
        data_files = []
        for file_name in extracted_files:
            full_path = os.path.join(temp_dir, file_name)
            if Path(file_name).suffix.lower() in ['.xlsx', '.xls', '.csv']:
                data_files.append(full_path)
        
        # İlk uygun dosyayı import et
        if data_files:
            result = AdvancedFileUtils._import_excel_advanced(data_files[0])
            result['metadata']['zip_contents'] = extracted_files
            result['file_info']['format'] = 'ZIP Archive'
            return result
        
        return {'success': False, 'message': 'ZIP içinde uygun veri dosyası bulunamadı'}
    
    @staticmethod
    def export_data_advanced(data, export_options=None):
        """
        Gelişmiş veri export işlemi
        
        Args:
            data: Export edilecek veri (DataFrame veya dict)
            export_options: Export seçenekleri
            
        Returns:
            dict: Export sonucu
        """
        if export_options is None:
            export_options = AdvancedFileUtils._get_export_options()
        
        if not export_options:
            return {'success': False, 'message': 'Export iptal edildi'}
        
        try:
            format_type = export_options['format']
            
            if format_type == 'excel':
                return AdvancedFileUtils._export_excel_advanced(data, export_options)
            elif format_type == 'csv':
                return AdvancedFileUtils._export_csv_advanced(data, export_options)
            elif format_type == 'json':
                return AdvancedFileUtils._export_json(data, export_options)
            elif format_type == 'xml':
                return AdvancedFileUtils._export_xml(data, export_options)
            elif format_type == 'word':
                return AdvancedFileUtils._export_word(data, export_options)
            elif format_type == 'zip':
                return AdvancedFileUtils._export_zip(data, export_options)
            else:
                return {'success': False, 'message': f'Desteklenmeyen format: {format_type}'}
                
        except Exception as e:
            return {
                'success': False,
                'message': f'Export hatası: {str(e)}',
                'error': str(e)
            }
    
    @staticmethod
    def _get_export_options():
        """Export seçenekleri dialog'u"""
        # Basit format seçimi (daha gelişmiş dialog eklenebilir)
        filetypes = [
            ("Excel Dosyası", "*.xlsx"),
            ("CSV Dosyası", "*.csv"),
            ("JSON Dosyası", "*.json"),
            ("XML Dosyası", "*.xml")
        ]
        
        if PYTHON_DOCX_AVAILABLE:
            filetypes.append(("Word Dosyası", "*.docx"))
        
        filetypes.append(("ZIP Arşivi", "*.zip"))
        
        file_path = filedialog.asksaveasfilename(
            title="Gelişmiş Veri Export",
            filetypes=filetypes,
            defaultextension=".xlsx"
        )
        
        if not file_path:
            return None
        
        # Format'ı dosya uzantısından belirle
        ext = Path(file_path).suffix.lower()
        format_map = {
            '.xlsx': 'excel',
            '.csv': 'csv',
            '.json': 'json',
            '.xml': 'xml',
            '.docx': 'word',
            '.zip': 'zip'
        }
        
        return {
            'path': file_path,
            'format': format_map.get(ext, 'excel'),
            'include_metadata': True,
            'include_timestamp': True
        }
    
    @staticmethod
    def _export_excel_advanced(data, options):
        """Gelişmiş Excel export"""
        with pd.ExcelWriter(options['path'], engine='openpyxl') as writer:
            # Ana veri
            if isinstance(data, pd.DataFrame):
                data.to_excel(writer, sheet_name='Spektrum_Verileri', index=False)
            elif isinstance(data, dict):
                for sheet_name, df in data.items():
                    if isinstance(df, pd.DataFrame):
                        df.to_excel(writer, sheet_name=sheet_name[:31], index=False)
            
            # Metadata sayfası
            if options.get('include_metadata', True):
                metadata_df = pd.DataFrame({
                    'Özellik': ['Export Tarihi', 'Format', 'Sürüm', 'Kaynak'],
                    'Değer': [
                        datetime.now().strftime('%Y-%m-%d %H:%M:%S'),
                        'Excel (.xlsx)',
                        'TBDY-2018 v1.0',
                        'Modüler Spektrum Analizi'
                    ]
                })
                metadata_df.to_excel(writer, sheet_name='Metadata', index=False)
        
        return {
            'success': True,
            'message': f'Excel dosyası başarıyla kaydedildi: {options["path"]}',
            'path': options['path']
        }
    
    @staticmethod
    def _export_csv_advanced(data, options):
        """Gelişmiş CSV export"""
        if isinstance(data, pd.DataFrame):
            data.to_csv(options['path'], index=False, encoding='utf-8-sig')
        elif isinstance(data, dict) and len(data) > 0:
            # İlk DataFrame'i kaydet
            first_df = next(iter(data.values()))
            if isinstance(first_df, pd.DataFrame):
                first_df.to_csv(options['path'], index=False, encoding='utf-8-sig')
        
        return {
            'success': True,
            'message': f'CSV dosyası başarıyla kaydedildi: {options["path"]}',
            'path': options['path']
        }
    
    @staticmethod
    def _export_json(data, options):
        """JSON export"""
        export_data = {}
        
        if isinstance(data, pd.DataFrame):
            export_data['data'] = data.to_dict('records')
        elif isinstance(data, dict):
            for key, value in data.items():
                if isinstance(value, pd.DataFrame):
                    export_data[key] = value.to_dict('records')
                else:
                    export_data[key] = value
        
        if options.get('include_metadata', True):
            export_data['metadata'] = {
                'export_date': datetime.now().isoformat(),
                'format': 'JSON',
                'version': 'TBDY-2018 v1.0',
                'source': 'Modüler Spektrum Analizi'
            }
        
        with open(options['path'], 'w', encoding='utf-8') as f:
            json.dump(export_data, f, indent=2, ensure_ascii=False)
        
        return {
            'success': True,
            'message': f'JSON dosyası başarıyla kaydedildi: {options["path"]}',
            'path': options['path']
        }
    
    @staticmethod
    def _export_xml(data, options):
        """XML export"""
        root = ET.Element("TBDYSpektrum")
        
        # Metadata
        if options.get('include_metadata', True):
            metadata = ET.SubElement(root, "Metadata")
            ET.SubElement(metadata, "ExportDate").text = datetime.now().isoformat()
            ET.SubElement(metadata, "Format").text = "XML"
            ET.SubElement(metadata, "Version").text = "TBDY-2018 v1.0"
        
        # Ana veri
        if isinstance(data, pd.DataFrame):
            data_elem = ET.SubElement(root, "Data")
            for _, row in data.iterrows():
                row_elem = ET.SubElement(data_elem, "Row")
                for col, value in row.items():
                    ET.SubElement(row_elem, str(col).replace(' ', '_')).text = str(value)
        
        # XML'i kaydet
        tree = ET.ElementTree(root)
        tree.write(options['path'], encoding='utf-8', xml_declaration=True)
        
        return {
            'success': True,
            'message': f'XML dosyası başarıyla kaydedildi: {options["path"]}',
            'path': options['path']
        }
    
    @staticmethod
    def _export_word(data, options):
        """Word dokümanı export"""
        if not PYTHON_DOCX_AVAILABLE:
            return {'success': False, 'message': 'python-docx kütüphanesi gerekli'}
        
        doc = Document()
        doc.add_heading('TBDY-2018 Spektrum Analizi Raporu', 0)
        
        # Tarih ve saat
        if options.get('include_timestamp', True):
            doc.add_paragraph(f'Oluşturulma Tarihi: {datetime.now().strftime("%Y-%m-%d %H:%M:%S")}')
        
        # Veri tablosu
        if isinstance(data, pd.DataFrame):
            doc.add_heading('Spektrum Verileri', level=1)
            
            # Tablo oluştur
            table = doc.add_table(rows=1, cols=len(data.columns))
            table.style = 'Table Grid'
            
            # Header
            header_cells = table.rows[0].cells
            for i, col in enumerate(data.columns):
                header_cells[i].text = str(col)
            
            # Veri satırları (ilk 100 satır)
            for _, row in data.head(100).iterrows():
                row_cells = table.add_row().cells
                for i, value in enumerate(row):
                    row_cells[i].text = str(value)
        
        doc.save(options['path'])
        
        return {
            'success': True,
            'message': f'Word dokümanı başarıyla kaydedildi: {options["path"]}',
            'path': options['path']
        }
    
    @staticmethod
    def _export_zip(data, options):
        """ZIP arşivi export"""
        zip_path = options['path']
        temp_dir = tempfile.mkdtemp()
        
        with zipfile.ZipFile(zip_path, 'w', zipfile.ZIP_DEFLATED) as zipf:
            # Excel dosyası ekle
            excel_path = os.path.join(temp_dir, 'spektrum_data.xlsx')
            excel_result = AdvancedFileUtils._export_excel_advanced(data, {'path': excel_path})
            if excel_result['success']:
                zipf.write(excel_path, 'spektrum_data.xlsx')
            
            # JSON dosyası ekle
            json_path = os.path.join(temp_dir, 'spektrum_data.json')
            json_result = AdvancedFileUtils._export_json(data, {'path': json_path, 'include_metadata': True})
            if json_result['success']:
                zipf.write(json_path, 'spektrum_data.json')
        
        return {
            'success': True,
            'message': f'ZIP arşivi başarıyla kaydedildi: {options["path"]}',
            'path': options['path']
        } 